"""
DocDrafterAgent — drafts a Confidential Information Memorandum (CIM) as a DOCX.

Sections:
  1. Executive Summary
  2. Business Description
  3. Management Team
  4. Financial Overview
  5. Market Opportunity
"""
from __future__ import annotations

import logging
import os
import re
from datetime import datetime
from pathlib import Path

from agents.base import BaseAgent
from agents.prompt_builder import PromptBuilder
from engine.llm import ask_llm

logger = logging.getLogger(__name__)

_OUTPUT_DIR = str(Path(__file__).resolve().parent.parent.parent.parent / "data" / "outputs")

_CIM_SECTIONS = [
    ("executive_summary", "Executive Summary"),
    ("business_description", "Business Description"),
    ("management", "Management Team"),
    ("financials", "Financial Overview"),
    ("market", "Market Opportunity"),
]


class DocDrafterAgent(BaseAgent):
    def __init__(self, deal_id: str, input_payload: dict):
        super().__init__(
            agent_type="doc_drafter",
            task_name="cim_draft",
            deal_id=deal_id,
            input_payload=input_payload,
        )
        self.system_prompt = PromptBuilder.get_system_prompt("doc_drafter")

    def run(self) -> str:
        try:
            self.think("Drafting Confidential Information Memorandum (CIM).")
            doc_context = self._extract_document_context()
            deal_info = self._get_deal_info()
            dcf_result = self._get_latest_dcf_output()
            deal_name = deal_info.get("deal_name", "Deal")
            company = deal_info.get("company_name", "Company")

            sections: dict[str, str] = {}
            for section_key, section_label in _CIM_SECTIONS:
                self.act("ask_llm", f"drafting CIM section: {section_label}")
                prompt = PromptBuilder.build_cim_section_prompt(
                    deal_info, doc_context, dcf_result, section_key
                )
                raw = ask_llm(self.system_prompt, prompt)
                sections[section_key] = raw.strip()
                self.observe(f"Section '{section_label}' drafted ({len(sections[section_key])} chars).")

            os.makedirs(_OUTPUT_DIR, exist_ok=True)
            date_str = datetime.now().strftime("%Y%m%d")
            safe_name = re.sub(r"[^\w\-]", "_", deal_name)
            docx_filename = f"{safe_name}_CIM_{date_str}.docx"
            docx_path = os.path.join(_OUTPUT_DIR, docx_filename)

            self.act("docx_writer", "writing CIM to DOCX")
            self._write_docx(docx_path, company, deal_info, sections)
            self.observe(f"CIM DOCX written: {docx_filename}")

            self._register_output(docx_path, output_type="docx", output_category="cim")
            self.complete(confidence=0.80)

        except Exception as exc:
            logger.exception("DocDrafterAgent failed for deal %s", self.deal_id)
            self.fail(str(exc))

        return self.run_id

    def _write_docx(
        self,
        docx_path: str,
        company: str,
        deal_info: dict,
        sections: dict[str, str],
    ) -> None:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ---- Document styles ----
        # Title
        title_para = doc.add_heading(company, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.runs[0]
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # Subtitle
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run("CONFIDENTIAL INFORMATION MEMORANDUM")
        sub_run.bold = True
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"{deal_info.get('deal_type', 'Transaction')} | {deal_info.get('industry', '')} | "
            f"{datetime.now().strftime('%B %Y')}"
        )
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.add_paragraph()  # spacer

        # Disclaimer
        disc_para = doc.add_paragraph()
        disc_run = disc_para.add_run(
            "CONFIDENTIAL — This document is intended solely for the named recipient. "
            "It may not be reproduced or distributed without prior written consent. "
            "This document does not constitute an offer or solicitation."
        )
        disc_run.font.size = Pt(8)
        disc_run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
        disc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # ---- Sections ----
        for section_key, section_label in _CIM_SECTIONS:
            content = sections.get(section_key, "")
            if not content:
                continue

            # Section heading
            heading = doc.add_heading(section_label, level=1)
            heading_run = heading.runs[0]
            heading_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

            # Section body — split by double newlines for paragraphs
            for para_text in content.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                p = doc.add_paragraph()
                p.add_run(para_text).font.size = Pt(10)

            doc.add_paragraph()  # spacer between sections

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        doc.save(docx_path)
